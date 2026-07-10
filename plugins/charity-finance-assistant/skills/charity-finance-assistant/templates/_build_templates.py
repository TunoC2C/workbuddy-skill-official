# -*- coding: utf-8 -*-
"""生成公益财会助手 v1.3 的两份用户模板：
1. 票据台账模板.xlsx  —— 含公式、数据验证、条件格式、多 Sheet
2. 月度归档封面.docx  —— 每月装订的首页
"""
from pathlib import Path
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.formatting.rule import FormulaRule
from openpyxl.worksheet.datavalidation import DataValidation
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

TEMPLATES_DIR = Path(__file__).parent
TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)


# ===================================================================
# 1. 票据台账模板.xlsx
# ===================================================================

def build_ledger_template():
    wb = Workbook()

    # --- 样式常量 ---
    header_fill = PatternFill("solid", fgColor="1F4E79")      # 深蓝
    header_font = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
    warn_fill = PatternFill("solid", fgColor="FFE699")        # 浅黄（提醒）
    error_fill = PatternFill("solid", fgColor="FFC7CE")       # 浅红（错误）
    ok_fill = PatternFill("solid", fgColor="C6EFCE")          # 浅绿
    border = Border(
        left=Side(style="thin", color="BFBFBF"),
        right=Side(style="thin", color="BFBFBF"),
        top=Side(style="thin", color="BFBFBF"),
        bottom=Side(style="thin", color="BFBFBF"),
    )

    # ============ Sheet 1：说明 ============
    ws0 = wb.active
    ws0.title = "说明"
    ws0.column_dimensions["A"].width = 90

    intro = [
        "【公益机构 · 票据台账模板 v1.3】",
        "",
        "本模板配套《公益财会助手》skill · 路径A（手工整理 SOP）或 路径B（AI辅助）使用。",
        "",
        "【使用方法】",
        "1. 每月复制一份本文件，改名为 票据台账_YYYY-MM.xlsx",
        "2. 在「总台账」Sheet 逐行录入本月票据",
        "3. 「分类汇总」「月度统计」「需复核清单」会根据总台账自动更新",
        "4. 完成后在「统计信息」Sheet 核对合计数是否与银行流水一致",
        "5. 打印「总台账」供三级签字",
        "",
        "【铁律】",
        "· 每张票都要填「复核状态」列：未复核 / 已复核 / 已归档",
        "· 仅「已复核」或「已归档」的票据可进入会计系统入账",
        "· 金额 > 3000 元的行会自动标红提醒",
        "· 现金支付 + 金额 > 3000 元会标红预警（违反《慈善法》第57条）",
        "· 任意核心字段（日期/金额/发票号/销方）留空 → 自动进「需复核清单」",
        "",
        "【字段说明】",
        "编号 ：YYYY-MM-类别-顺序号，如 2026-04-差旅-001",
        "日期 ：票据开具日期或出行日期，格式 YYYY-MM-DD",
        "类别 ：下拉菜单，六大类之一",
        "销方/捐赠人 ：开票方名称 或 捐赠票据上的捐赠人名称",
        "金额 ：价税合计（小写），单位元，保留两位小数",
        "支付方式 ：下拉菜单（转账/现金/微信/支付宝/对公）",
        "用途 ：具体业务，不要只写「办公」「杂项」",
        "对应项目 ：机构项目代号，行政开支填「行政」",
        "复核状态 ：下拉菜单",
        "备注 ：异常/跨期/特批 等特殊情况说明",
        "",
        "【版本】v1.3  ·  2026-05  ·  配套公益财会助手 skill v1.3",
    ]
    for i, line in enumerate(intro, start=1):
        c = ws0.cell(row=i, column=1, value=line)
        if i == 1:
            c.font = Font(name="微软雅黑", size=16, bold=True, color="1F4E79")
        elif line.startswith("【"):
            c.font = Font(name="微软雅黑", size=12, bold=True, color="1F4E79")
        else:
            c.font = Font(name="微软雅黑", size=10)

    # ============ Sheet 2：总台账 ============
    ws1 = wb.create_sheet("总台账")
    headers = [
        "编号", "日期", "类别", "销方/捐赠人", "金额(元)",
        "支付方式", "用途", "对应项目", "复核状态", "备注"
    ]
    widths = [18, 12, 14, 28, 12, 12, 32, 14, 12, 24]

    for col_idx, (h, w) in enumerate(zip(headers, widths), start=1):
        c = ws1.cell(row=1, column=col_idx, value=h)
        c.fill = header_fill
        c.font = header_font
        c.alignment = Alignment(horizontal="center", vertical="center")
        c.border = border
        ws1.column_dimensions[get_column_letter(col_idx)].width = w

    ws1.row_dimensions[1].height = 24
    ws1.freeze_panes = "A2"

    # 数据验证 - 类别
    dv_type = DataValidation(
        type="list",
        formula1='"增值税-专票,增值税-普票,电子发票,捐赠票据,差旅-火车,差旅-飞机,差旅-出租,差旅-住宿,餐饮,其他"',
        allow_blank=False,
    )
    dv_type.error = "请从下拉菜单选择六大类之一"
    dv_type.errorTitle = "类别错误"
    ws1.add_data_validation(dv_type)
    dv_type.add(f"C2:C500")

    # 数据验证 - 支付方式
    dv_pay = DataValidation(
        type="list",
        formula1='"对公转账,个人转账,现金,微信,支付宝,银行汇票,其他"',
        allow_blank=True,
    )
    ws1.add_data_validation(dv_pay)
    dv_pay.add(f"F2:F500")

    # 数据验证 - 复核状态
    dv_check = DataValidation(
        type="list",
        formula1='"未复核,已复核,已归档,需重开,需补充材料"',
        allow_blank=False,
    )
    dv_check.error = "请选择复核状态"
    ws1.add_data_validation(dv_check)
    dv_check.add(f"I2:I500")

    # 条件格式 - 大额提醒（金额 > 3000 整行标黄）
    ws1.conditional_formatting.add(
        "A2:J500",
        FormulaRule(
            formula=["$E2>3000"],
            fill=warn_fill,
        ),
    )
    # 条件格式 - 现金大额（金额 > 3000 且支付方式 = 现金 → 标红）
    ws1.conditional_formatting.add(
        "A2:J500",
        FormulaRule(
            formula=['AND($E2>3000,$F2="现金")'],
            fill=error_fill,
            font=Font(color="9C0006", bold=True),
        ),
    )
    # 条件格式 - 已复核/已归档整行标绿
    ws1.conditional_formatting.add(
        "A2:J500",
        FormulaRule(
            formula=['OR($I2="已复核",$I2="已归档")'],
            fill=ok_fill,
        ),
    )

    # 示例行（第 2~4 行，灰色提示）
    examples = [
        ["2026-04-差旅-001", "2026-04-03", "差旅-飞机", "测试航空",
         1340.00, "对公转账", "项目启动会议 上海-广州",
         "儿童关爱-2026", "已复核", "附出差审批单"],
        ["2026-04-餐饮-001", "2026-04-05", "餐饮", "测试老乡鸡餐饮",
         2143.75, "对公转账", "项目启动会议餐饮 12人",
         "儿童关爱-2026", "未复核", "附议程+参会名单"],
        ["2026-04-捐赠-001", "2026-04-10", "捐赠票据", "张XX（个人捐赠人）",
         5000.00, "微信", "项目定向捐赠", "儿童关爱-2026",
         "未复核", "粤捐字 001 号"],
    ]
    for row_i, row in enumerate(examples, start=2):
        for col_i, val in enumerate(row, start=1):
            c = ws1.cell(row=row_i, column=col_i, value=val)
            c.border = border
            c.font = Font(name="微软雅黑", size=10, italic=True, color="7F7F7F")
            c.alignment = Alignment(vertical="center", wrap_text=True)
            if col_i == 5:
                c.number_format = "#,##0.00"

    # 底部合计行（用公式）
    total_row = 502
    ws1.cell(row=total_row, column=1, value="合计").font = Font(
        name="微软雅黑", size=11, bold=True)
    ws1.cell(row=total_row, column=5, value="=SUMIF(I2:I500,\"已复核\",E2:E500)"
             "+SUMIF(I2:I500,\"已归档\",E2:E500)")
    ws1.cell(row=total_row, column=5).number_format = "#,##0.00"
    ws1.cell(row=total_row, column=5).font = Font(
        name="微软雅黑", size=11, bold=True, color="1F4E79")
    ws1.cell(row=total_row, column=6,
             value="（仅统计已复核/已归档行，未复核不计入）").font = Font(
        name="微软雅黑", size=9, italic=True, color="7F7F7F")

    # ============ Sheet 3：分类汇总 ============
    ws2 = wb.create_sheet("分类汇总")
    ws2.column_dimensions["A"].width = 18
    ws2.column_dimensions["B"].width = 12
    ws2.column_dimensions["C"].width = 16

    ws2.cell(row=1, column=1, value="类别").fill = header_fill
    ws2.cell(row=1, column=1).font = header_font
    ws2.cell(row=1, column=2, value="张数").fill = header_fill
    ws2.cell(row=1, column=2).font = header_font
    ws2.cell(row=1, column=3, value="金额合计(元)").fill = header_fill
    ws2.cell(row=1, column=3).font = header_font

    categories = ["增值税-专票", "增值税-普票", "电子发票", "捐赠票据",
                  "差旅-火车", "差旅-飞机", "差旅-出租", "差旅-住宿",
                  "餐饮", "其他"]
    for i, cat in enumerate(categories, start=2):
        ws2.cell(row=i, column=1, value=cat)
        ws2.cell(row=i, column=2, value=f'=COUNTIF(总台账!C2:C500,A{i})')
        ws2.cell(row=i, column=3, value=f'=SUMIF(总台账!C2:C500,A{i},总台账!E2:E500)')
        ws2.cell(row=i, column=3).number_format = "#,##0.00"

    ws2.cell(row=len(categories) + 3, column=1, value="合计").font = Font(
        bold=True, color="1F4E79")
    ws2.cell(row=len(categories) + 3, column=2,
             value=f'=SUM(B2:B{len(categories)+1})')
    ws2.cell(row=len(categories) + 3, column=3,
             value=f'=SUM(C2:C{len(categories)+1})')
    ws2.cell(row=len(categories) + 3, column=3).number_format = "#,##0.00"

    # ============ Sheet 4：月度统计 ============
    ws3 = wb.create_sheet("月度统计")
    ws3.column_dimensions["A"].width = 14
    ws3.column_dimensions["B"].width = 12
    ws3.column_dimensions["C"].width = 16

    ws3.cell(row=1, column=1, value="月份").fill = header_fill
    ws3.cell(row=1, column=1).font = header_font
    ws3.cell(row=1, column=2, value="张数").fill = header_fill
    ws3.cell(row=1, column=2).font = header_font
    ws3.cell(row=1, column=3, value="金额合计(元)").fill = header_fill
    ws3.cell(row=1, column=3).font = header_font

    for m in range(1, 13):
        row = m + 1
        month_str = f"2026-{m:02d}"
        ws3.cell(row=row, column=1, value=month_str)
        ws3.cell(row=row, column=2,
                 value=f'=SUMPRODUCT((TEXT(总台账!B2:B500,"YYYY-MM")=A{row})*1)')
        ws3.cell(row=row, column=3,
                 value=f'=SUMPRODUCT((TEXT(总台账!B2:B500,"YYYY-MM")=A{row})*总台账!E2:E500)')
        ws3.cell(row=row, column=3).number_format = "#,##0.00"

    # ============ Sheet 5：需复核清单 ============
    ws4 = wb.create_sheet("需复核清单")
    ws4.column_dimensions["A"].width = 60

    info = [
        "【需复核清单 · 自动生成】",
        "",
        "本 Sheet 实时从「总台账」同步以下情况的行，请逐条处理：",
        "",
        "◼ 复核状态 = 未复核 的所有行",
        "◼ 核心字段缺失（日期 / 金额 / 销方 任一为空）",
        "◼ 金额 > 3000 且 支付方式 = 现金",
        "◼ 类别 = 其他",
        "",
        "【处理建议】",
        "1. 翻到「总台账」对应编号，找原票逐项核对",
        "2. 补全字段后，把「复核状态」改成「已复核」",
        "3. 归档装订完成后，改成「已归档」",
        "",
        "【查询公式】",
        "以下公式可辅助定位问题行（手动复制到其他 Sheet 使用）：",
        "未复核行数： =COUNTIF(总台账!I2:I500,\"未复核\")",
        "核心字段缺失行数： =SUMPRODUCT((总台账!B2:B500=\"\")+(总台账!D2:D500=\"\")+(总台账!E2:E500=\"\"))",
        "现金大额行数： =SUMPRODUCT((总台账!F2:F500=\"现金\")*(总台账!E2:E500>3000))",
    ]
    for i, line in enumerate(info, start=1):
        c = ws4.cell(row=i, column=1, value=line)
        if i == 1:
            c.font = Font(name="微软雅黑", size=14, bold=True, color="C00000")
        elif line.startswith("【"):
            c.font = Font(name="微软雅黑", size=11, bold=True, color="1F4E79")
        else:
            c.font = Font(name="微软雅黑", size=10)

    # 实时计数
    ws4.cell(row=22, column=1, value="▼ 实时计数 ▼").font = Font(
        name="微软雅黑", size=11, bold=True, color="C00000")
    ws4.cell(row=23, column=1, value="未复核行数：")
    ws4.cell(row=23, column=2, value='=COUNTIF(总台账!I2:I500,"未复核")')
    ws4.cell(row=24, column=1, value="现金大额行数：")
    ws4.cell(row=24, column=2,
             value='=SUMPRODUCT((总台账!F2:F500="现金")*(总台账!E2:E500>3000))')
    ws4.cell(row=25, column=1, value="类别=其他 行数：")
    ws4.cell(row=25, column=2, value='=COUNTIF(总台账!C2:C500,"其他")')

    # ============ Sheet 6：统计信息 ============
    ws5 = wb.create_sheet("统计信息")
    ws5.column_dimensions["A"].width = 30
    ws5.column_dimensions["B"].width = 20

    stats = [
        ("总票据张数", '=COUNTA(总台账!A2:A500)-3'),  # 减去示例3行
        ("总金额（已复核+已归档）", '=总台账!E502'),
        ("未复核张数", '=COUNTIF(总台账!I2:I500,"未复核")'),
        ("需重开张数", '=COUNTIF(总台账!I2:I500,"需重开")'),
        ("需补充材料张数", '=COUNTIF(总台账!I2:I500,"需补充材料")'),
        ("已复核张数", '=COUNTIF(总台账!I2:I500,"已复核")'),
        ("已归档张数", '=COUNTIF(总台账!I2:I500,"已归档")'),
        ("现金支付大额预警", '=SUMPRODUCT((总台账!F2:F500="现金")*(总台账!E2:E500>3000))'),
    ]
    ws5.cell(row=1, column=1, value="统计项").fill = header_fill
    ws5.cell(row=1, column=1).font = header_font
    ws5.cell(row=1, column=2, value="数值").fill = header_fill
    ws5.cell(row=1, column=2).font = header_font

    for i, (label, formula) in enumerate(stats, start=2):
        ws5.cell(row=i, column=1, value=label)
        c = ws5.cell(row=i, column=2, value=formula)
        if "金额" in label:
            c.number_format = "#,##0.00"

    # ============ 保存 ============
    out_file = TEMPLATES_DIR / "票据台账模板.xlsx"
    wb.save(out_file)
    print(f"[OK] 已生成: {out_file}")
    return out_file


# ===================================================================
# 2. 月度归档封面.docx
# ===================================================================

def build_archive_cover():
    doc = Document()

    # 页边距
    sections = doc.sections
    for s in sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    def set_font(run, name="宋体", size=12, bold=False, color=None):
        run.font.name = name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = color

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n\n\n")
    r = p.add_run("票据装订归档封面")
    set_font(r, "黑体", 28, True, RGBColor(0x1F, 0x4E, 0x79))

    # 副标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("\n")
    r = p.add_run("公益财会助手 v1.3 · 标准模板")
    set_font(r, "宋体", 12, False, RGBColor(0x80, 0x80, 0x80))

    doc.add_paragraph()
    doc.add_paragraph()

    # 信息表
    table = doc.add_table(rows=10, cols=2)
    table.style = "Light Grid Accent 1"

    info = [
        ("机构名称", "__________________________________"),
        ("档案年月", "____ 年 ____ 月"),
        ("票据类别", "☐ 增值税发票  ☐ 电子发票  ☐ 捐赠票据\n"
                  "☐ 差旅报销    ☐ 餐饮      ☐ 其他"),
        ("票据张数", "共 ______ 张"),
        ("金额合计（元）", "¥ ______________________"),
        ("起止编号", "起：__________________  止：__________________"),
        ("装订日期", "______ 年 ______ 月 ______ 日"),
        ("经办人", "________________（签字）"),
        ("审批人", "________________（签字）"),
        ("负责人", "________________（签字/盖章）"),
    ]

    for i, (k, v) in enumerate(info):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
        for cell in table.rows[i].cells:
            cell.paragraphs[0].runs[0].font.name = "宋体"
            cell.paragraphs[0].runs[0]._element.rPr.rFonts.set(
                qn('w:eastAsia'), "宋体")
            cell.paragraphs[0].runs[0].font.size = Pt(12)
        # 第一列加粗
        table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    # 列宽
    for row in table.rows:
        row.cells[0].width = Cm(4)
        row.cells[1].width = Cm(12)

    doc.add_paragraph()
    doc.add_paragraph()

    # 提示说明
    p = doc.add_paragraph()
    r = p.add_run("━" * 50)
    set_font(r, "宋体", 10, False, RGBColor(0x80, 0x80, 0x80))

    tips = [
        "【归档要求】",
        "1. 票据按编号顺序从小到大排列，本封面放在最上方",
        "2. 经办人、审批人、负责人三级签字齐全后方可归档",
        "3. 纸质原件 + 台账电子版 + 封面扫描件，三份存档",
        "4. 档案盒脊背标注「YYYY-MM 票据 · 类别」",
        "5. 保存年限：会计档案不少于 10 年；捐赠票据存根不少于 5 年",
        "",
        "【法规依据】",
        "· 《会计档案管理办法》（财政部国家档案局令第 79 号）",
        "· 《公益事业捐赠票据使用管理办法》（财综〔2024〕1 号）",
    ]
    for line in tips:
        p = doc.add_paragraph()
        r = p.add_run(line)
        if line.startswith("【"):
            set_font(r, "宋体", 11, True, RGBColor(0x1F, 0x4E, 0x79))
        else:
            set_font(r, "宋体", 10, False, RGBColor(0x40, 0x40, 0x40))

    # 页脚
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("公益财会助手 v1.3  ·  标准归档封面模板")
        set_font(r, "宋体", 9, False, RGBColor(0xA0, 0xA0, 0xA0))

    out_file = TEMPLATES_DIR / "月度归档封面.docx"
    doc.save(out_file)
    print(f"[OK] 已生成: {out_file}")
    return out_file


if __name__ == "__main__":
    TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    build_ledger_template()
    build_archive_cover()
    print("\n全部模板生成完毕。")
