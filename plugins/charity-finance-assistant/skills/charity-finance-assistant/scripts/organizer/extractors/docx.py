"""
v1.3 DOCX 抽取器：拆解原 extract_docx_fields（圈复杂度 23 → ≤ 5）。

TODO（留给后续迭代，对应工蜂 AI CR 建议 #3 拒绝原因）：
当前 _read_paragraphs() + _read_tables() 先后拼接，丢失了文档原始顺序。
对发票/报销单这类“表单型”文档不问题（字段顺序不敏感），
但如果未来需要提取“项目订立合同/按顺序重要的多段混合文档”，
需手动解析 doc.element.body 的 XML 子元素（w:p 与 w:tbl）。
参考实现：https://python-docx.readthedocs.io/en/latest/dev/analysis/features/text/paragraph.html
"""
from organizer.extractors.base import empty_fields, fill_fields_from_text

try:
    import docx as _docx_mod  # python-docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _read_paragraphs(doc):
    """读取所有非空段落"""
    return [p.text for p in doc.paragraphs if p.text.strip()]


def _read_tables(doc):
    """把所有表格按行拼接成单行字符串列表"""
    rows_text = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells)
            if row_text:
                rows_text.append(row_text)
    return rows_text


def _read_docx_text(docx_path):
    """
    读取 docx 全部段落+表格文本。
    返回 (text, error_msg)。
    """
    try:
        doc = _docx_mod.Document(str(docx_path))
    except Exception as e:
        return "", f"解析失败: {e}"
    parts = _read_paragraphs(doc) + _read_tables(doc)
    return "\n".join(parts), None


def extract_docx_fields(docx_path):
    """
    从 DOCX 提取发票要素。
    圈复杂度 ≈ 4：HAS_DOCX 检查 + 读文本失败检查 + 正常路径。
    """
    if not HAS_DOCX:
        return empty_fields("缺少python-docx")

    text, err = _read_docx_text(docx_path)
    if err:
        return empty_fields(err)

    fields = empty_fields("成功")
    fields["_raw_text"] = text
    fill_fields_from_text(fields, text)
    return fields
